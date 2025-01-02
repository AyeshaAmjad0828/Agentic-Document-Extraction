import pdfplumber
from PyPDF2 import PdfWriter, PdfReader
import os
from PIL import Image
from docx import Document
import os

def split_pdf_by_pages(input_pdf_path, output_folder):
    """
    Split a PDF document by pages, saving each page as a separate PDF file.

    Args:
        input_pdf_path (str): Path to the input PDF file.
        output_folder (str): Directory where the split pages will be saved.

    Returns:
        list of str: Paths to the individual page PDFs.
    """
    # Ensure the output folder exists
    os.makedirs(output_folder, exist_ok=True)
    
    # Initialize an empty list to hold paths of split page PDFs
    page_paths = []
    
    # Open the PDF document
    with pdfplumber.open(input_pdf_path) as pdf:
        for page_num in range(len(pdf.pages)):
            # Extract each page as a separate PDF
            page = pdf.pages[page_num]
            page_text = page.extract_text()
            
            # Create a writer object to save the page as a new PDF
            writer = PdfWriter()
            writer.add_page(PdfReader(input_pdf_path).pages[page_num])
            
            # Define path for the new single-page PDF file
            output_pdf_path = os.path.join(output_folder, f"page_{page_num + 1}.pdf")
            
            # Write the page to a separate PDF file
            with open(output_pdf_path, "wb") as output_pdf:
                writer.write(output_pdf)
            
            # Append the page path to the list
            page_paths.append(output_pdf_path)
            
            # Print page content (for demonstration; you might store this in a data structure instead)
            print(f"Page {page_num + 1} Content:\n", page_text)
    
    return page_paths



def split_images(input_folder, output_folder):
    """
    Splits an image document, treating each image file in a folder as a separate page.
    
    Args:
        input_folder (str): Path to the folder containing image files.
        output_folder (str): Directory to save the individual pages.
        
    Returns:
        list of str: Paths to individual page images.
    """
    os.makedirs(output_folder, exist_ok=True)
    page_paths = []
    
    for index, filename in enumerate(sorted(os.listdir(input_folder))):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif')):
            img = Image.open(os.path.join(input_folder, filename))
            output_path = os.path.join(output_folder, f"page_{index + 1}.png")
            img.save(output_path)
            page_paths.append(output_path)
            print(f"Saved image page: {output_path}")
    
    return page_paths

def split_word_document(input_docx_path, output_folder):
    """
    Splits a Word document by pages by creating separate documents for each page.
    
    Args:
        input_docx_path (str): Path to the input Word document.
        output_folder (str): Directory to save individual page documents.
        
    Returns:
        list of str: Paths to individual page Word documents.
    """
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(input_docx_path)
    page_docs = []
    
    # Split the document manually by paragraphs
    page_number = 1
    current_doc = Document()
    for paragraph in doc.paragraphs:
        current_doc.add_paragraph(paragraph.text)
        
        # Placeholder logic for splitting on some condition.
        # Here, we assume a new page starts after every 10 paragraphs.
        if len(current_doc.paragraphs) >= 10:  # Adjust this number as needed
            output_path = os.path.join(output_folder, f"page_{page_number}.docx")
            current_doc.save(output_path)
            page_docs.append(output_path)
            print(f"Saved Word page: {output_path}")
            
            # Prepare for the next page
            page_number += 1
            current_doc = Document()
    
    # Save any remaining paragraphs as the last page
    if current_doc.paragraphs:
        output_path = os.path.join(output_folder, f"page_{page_number}.docx")
        current_doc.save(output_path)
        page_docs.append(output_path)
        print(f"Saved Word page: {output_path}")
    
    return page_docs

# Example usage

input_image_folder = "path_to_images_folder"  # Folder containing individual images as "pages"
output_image_folder = "split_image_pages"
image_page_paths = split_images(input_image_folder, output_image_folder)
print("Image pages saved:", image_page_paths)


input_docx_path = "sample_document.docx"
output_word_folder = "split_word_pages"
word_page_paths = split_word_document(input_docx_path, output_word_folder)
print("Word document pages saved:", word_page_paths)



input_pdf_path = "MMCW878403-My Muscle Chef.pdf"  # Path to your multi-page PDF file
output_folder = "split_pages"  # Folder to save split pages
page_paths = split_pdf_by_pages(input_pdf_path, output_folder)
print("Pages saved individually as PDFs:", page_paths)
